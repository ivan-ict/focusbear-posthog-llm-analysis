"""Generate an aggregate-only supervisor report from the workbook output."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
import os
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from openpyxl import load_workbook


@dataclass(slots=True)
class SupervisorReportData:
    """Aggregate data extracted from the workbook for DOCX rendering."""

    generated_at: str
    cohort_id: str
    cohort_name: str
    cohort_users_available: int
    users_analyzed: int
    applied_user_limit: str
    lookback_days: int
    onboarding_completed: int
    onboarding_completed_pct: str
    categories: list[tuple[str, int, str]]
    dropoff_points: list[tuple[str, int]]
    blocking_stages: list[tuple[str, int, str]]
    error_event_totals: list[tuple[str, int, int, str]]
    error_breakdown: list[tuple[str, str, str, int, int, str]]
    key_findings: list[str]


def generate_supervisor_report(workbook_path: Path, output_path: Path) -> Path:
    """Build a professional DOCX report from an existing workbook output."""
    report_data = read_supervisor_report_data(workbook_path)

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Supervisor report generation requires python-docx. "
            "Install dependencies from requirements.txt first."
        ) from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "Focus Bear Onboarding Analysis Report"

    title = document.add_heading("Focus Bear Onboarding Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        f"Cohort {report_data.cohort_id}: {report_data.cohort_name or 'Unknown'}\n"
        f"Generated from local workbook output on {report_data.generated_at}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(_executive_summary(report_data))

    document.add_heading("Cohort and Method", level=1)
    document.add_paragraph(
        "This report was generated offline from the local onboarding workbook. "
        "It uses aggregate-only workbook metrics and does not include user identifiers, emails, or raw notes."
    )
    _add_table(
        document,
        headers=["Metric", "Value"],
        rows=[
            ("Cohort ID", report_data.cohort_id),
            ("Cohort Name", report_data.cohort_name or "Unknown"),
            ("Users Available", str(report_data.cohort_users_available)),
            ("Users Analyzed", str(report_data.users_analyzed)),
            ("Applied User Limit", report_data.applied_user_limit),
            ("Lookback Window (Days)", str(report_data.lookback_days)),
            ("Onboarding Completed", str(report_data.onboarding_completed)),
            ("Onboarding Completed %", report_data.onboarding_completed_pct),
        ],
    )

    document.add_heading("Key Findings", level=1)
    for finding in report_data.key_findings:
        document.add_paragraph(finding, style="List Bullet")

    document.add_heading("Journey Analysis", level=1)
    document.add_paragraph(_journey_analysis(report_data))
    _add_table(document, ["Category", "Count", "Percent"], report_data.categories)
    _add_table(document, ["Top Dropoff Point", "Count"], report_data.dropoff_points[:10])

    document.add_heading("Backend / Error Analysis", level=1)
    document.add_paragraph(_backend_analysis(report_data))
    _add_table(
        document,
        ["Event", "Raw Events", "Affected Users", "Affected Users %"],
        report_data.error_event_totals[:10],
    )
    _add_table(
        document,
        ["Event", "Endpoint URL", "Status Code", "Raw Events", "Affected Users", "Affected Users %"],
        report_data.error_breakdown[:10],
    )

    document.add_heading("Blocking Schedule Analysis", level=1)
    document.add_paragraph(_blocking_analysis(report_data))
    _add_table(document, ["Blocking Stage", "Count", "Percent"], report_data.blocking_stages)

    document.add_heading("Recommendations / Next Steps", level=1)
    for recommendation in _recommendations(report_data):
        document.add_paragraph(recommendation, style="List Bullet")

    document.add_heading("Appendix", level=1)
    document.add_paragraph(
        "Appendix tables summarize the current workbook output only. "
        "No row-level user identifiers or free-text notes are included in this report."
    )

    with TemporaryDirectory() as temp_dir:
        chart_paths = _create_chart_images(report_data, Path(temp_dir))
        for title_text, chart_path in chart_paths:
            document.add_heading(title_text, level=2)
            document.add_picture(str(chart_path), width=Inches(6.5))

    document.save(output_path)
    return output_path


def read_supervisor_report_data(workbook_path: Path) -> SupervisorReportData:
    """Extract the aggregate workbook sections used by the supervisor report."""
    workbook = load_workbook(workbook_path, data_only=True, read_only=True)
    worksheet = workbook["Summary"]
    rows = [tuple(row) for row in worksheet.iter_rows(values_only=True)]

    metrics_header = _find_row(rows, "Metric", "Value")
    categories_header = _find_row(rows, "Category", "Count")
    dropoff_header = _find_row(rows, "Top Dropoff Point", "Count")
    blocking_header = _find_row(rows, "Blocking Schedule Deepest Stage", "Count")
    error_totals_title = _find_row(rows, "Error Event Totals")
    error_breakdown_title = _find_row(rows, "Error Breakdown")
    findings_title = _find_row(rows, "Key Finding")

    metrics = _collect_key_value_rows(rows, metrics_header + 1)
    categories = _collect_table_rows(rows, categories_header + 1, 3)
    dropoff_points = _collect_table_rows(rows, dropoff_header + 1, 2)
    blocking_stages = _collect_table_rows(rows, blocking_header + 1, 3)
    error_event_totals = _collect_table_rows(rows, error_totals_title + 2, 4)
    error_breakdown = _collect_table_rows(rows, error_breakdown_title + 2, 6)
    key_findings = [str(row[0]) for row in _collect_table_rows(rows, findings_title + 1, 1)]

    return SupervisorReportData(
        generated_at=_format_metric_value(metrics.get("Generated At")),
        cohort_id=str(metrics.get("Cohort ID") or ""),
        cohort_name=str(metrics.get("Cohort Name") or ""),
        cohort_users_available=int(metrics.get("Cohort Users Available") or 0),
        users_analyzed=int(metrics.get("Users Analyzed") or 0),
        applied_user_limit=str(metrics.get("Applied User Limit") or ""),
        lookback_days=int(metrics.get("Lookback Window (Days)") or 0),
        onboarding_completed=int(metrics.get("Onboarding Completed") or 0),
        onboarding_completed_pct=str(metrics.get("Onboarding Completed %") or "0.0%"),
        categories=[_coerce_category_row(row) for row in categories],
        dropoff_points=[_coerce_dropoff_row(row) for row in dropoff_points],
        blocking_stages=[_coerce_blocking_row(row) for row in blocking_stages],
        error_event_totals=[_coerce_error_total_row(row) for row in error_event_totals],
        error_breakdown=[_coerce_error_breakdown_row(row) for row in error_breakdown],
        key_findings=key_findings,
    )


def _executive_summary(report_data: SupervisorReportData) -> str:
    """Return a deterministic executive summary paragraph."""
    top_category = report_data.categories[0] if report_data.categories else ("Unknown", 0, "0.0%")
    top_dropoff = report_data.dropoff_points[0] if report_data.dropoff_points else ("None", 0)
    return (
        f"This report reviews {report_data.users_analyzed} onboarding journeys from cohort "
        f"{report_data.cohort_id}. The current onboarding completion rate is "
        f"{report_data.onboarding_completed_pct} ({report_data.onboarding_completed} users). "
        f"The largest journey category is {top_category[0]} at {top_category[1]} users "
        f"({top_category[2]}). The most common dropoff point is {top_dropoff[0]} "
        f"with {top_dropoff[1]} users."
    )


def _journey_analysis(report_data: SupervisorReportData) -> str:
    """Return a short narrative for journey outcomes."""
    if not report_data.categories:
        return "No category data was available in the workbook summary."
    top_category = report_data.categories[0]
    return (
        f"Journey outcomes are dominated by {top_category[0]}, which affects "
        f"{top_category[1]} users. The dropoff table highlights where users most often stop "
        "before completing onboarding, which should guide the next investigation and experiment cycle."
    )


def _backend_analysis(report_data: SupervisorReportData) -> str:
    """Return a short narrative for backend and endpoint issues."""
    if not report_data.error_event_totals or report_data.error_event_totals[0][0] == "None":
        return "No canonical backend error events were recorded in the workbook summary."
    top_error = report_data.error_event_totals[0]
    top_breakdown = report_data.error_breakdown[0] if report_data.error_breakdown else None
    breakdown_text = ""
    if top_breakdown is not None:
        breakdown_text = (
            f" The most affected endpoint/status combination is {top_breakdown[0]} at "
            f"{top_breakdown[1]} with status {top_breakdown[2]}, affecting {top_breakdown[4]} users."
        )
    return (
        f"The leading canonical backend event is {top_error[0]}, with {top_error[1]} raw events "
        f"across {top_error[2]} affected users.{breakdown_text}"
    )


def _blocking_analysis(report_data: SupervisorReportData) -> str:
    """Return a short narrative for blocking schedule progression."""
    if not report_data.blocking_stages:
        return "No blocking schedule stage data was available in the workbook summary."
    top_stage = report_data.blocking_stages[0]
    return (
        f"The most common deepest blocking schedule stage is {top_stage[0]}, which covers "
        f"{top_stage[1]} users ({top_stage[2]}). This indicates how far users get into "
        "blocking setup before abandoning or completing the flow."
    )


def _recommendations(report_data: SupervisorReportData) -> list[str]:
    """Build deterministic recommendation bullets from the aggregate report data."""
    recommendations: list[str] = []
    if report_data.dropoff_points:
        recommendations.append(
            f"Prioritize investigation of the {report_data.dropoff_points[0][0]} step, "
            f"which is currently the most common dropoff point."
        )
    if report_data.error_breakdown:
        top_breakdown = report_data.error_breakdown[0]
        recommendations.append(
            f"Review the backend path for {top_breakdown[1]} with status {top_breakdown[2]}, "
            "because it is the most visible canonical backend issue in the current cohort."
        )
    recommendations.append(
        "Use the workbook category mix and blocking-stage progression to define the next onboarding experiment."
    )
    return recommendations


def _create_chart_images(report_data: SupervisorReportData, output_dir: Path) -> list[tuple[str, Path]]:
    """Generate local chart images for embedding in the DOCX report."""
    try:
        mpl_config_dir = output_dir / ".matplotlib"
        cache_dir = output_dir / ".cache"
        mpl_config_dir.mkdir(parents=True, exist_ok=True)
        cache_dir.mkdir(parents=True, exist_ok=True)
        os.environ.setdefault("MPLCONFIGDIR", str(mpl_config_dir))
        os.environ.setdefault("XDG_CACHE_HOME", str(cache_dir))
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Supervisor report generation requires matplotlib. "
            "Install dependencies from requirements.txt first."
        ) from exc

    charts: list[tuple[str, Path]] = []
    chart_specs = [
        (
            "Category Distribution",
            [row[0] for row in report_data.categories],
            [row[1] for row in report_data.categories],
            output_dir / "categories.png",
            "#1F4E79",
        ),
        (
            "Top Dropoff Points",
            [row[0] for row in report_data.dropoff_points[:10]],
            [row[1] for row in report_data.dropoff_points[:10]],
            output_dir / "dropoffs.png",
            "#C55A11",
        ),
        (
            "Canonical Error Events by Affected Users",
            [row[0] for row in report_data.error_event_totals[:10]],
            [row[2] for row in report_data.error_event_totals[:10]],
            output_dir / "errors.png",
            "#7F6000",
        ),
    ]

    for title, labels, values, chart_path, color in chart_specs:
        if not labels or not any(values):
            continue
        plt.figure(figsize=(8, 4.5))
        plt.barh(labels, values, color=color)
        plt.title(title)
        plt.tight_layout()
        plt.savefig(chart_path, dpi=180)
        plt.close()
        charts.append((title, chart_path))
    return charts


def _add_table(document: Any, headers: list[str], rows: list[tuple[Any, ...]]) -> None:
    """Add a simple grid table to the DOCX document."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def _find_row(rows: list[tuple[Any, ...]], first_cell: str, second_cell: str | None = None) -> int:
    """Return the index of the first matching row."""
    for index, row in enumerate(rows):
        first = str(row[0] or "").strip()
        second = str(row[1] or "").strip() if len(row) > 1 else ""
        if first != first_cell:
            continue
        if second_cell is not None and second != second_cell:
            continue
        return index
    raise ValueError(f"Could not find summary row starting with {first_cell!r}")


def _collect_key_value_rows(rows: list[tuple[Any, ...]], start_index: int) -> dict[str, Any]:
    """Collect summary metric rows until the next blank row."""
    output: dict[str, Any] = {}
    for row in rows[start_index:]:
        if row[0] is None:
            break
        output[str(row[0])] = row[1]
    return output


def _collect_table_rows(rows: list[tuple[Any, ...]], start_index: int, width: int) -> list[tuple[Any, ...]]:
    """Collect section rows until the next blank first cell."""
    output: list[tuple[Any, ...]] = []
    for row in rows[start_index:]:
        if row[0] is None:
            break
        output.append(tuple(row[:width]))
    return output


def _format_metric_value(value: Any) -> str:
    """Format workbook metric values for prose output."""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y %H:%M")
    return str(value or "")


def _coerce_category_row(row: tuple[Any, ...]) -> tuple[str, int, str]:
    return (str(row[0]), int(row[1] or 0), str(row[2] or "0.0%"))


def _coerce_dropoff_row(row: tuple[Any, ...]) -> tuple[str, int]:
    return (str(row[0]), int(row[1] or 0))


def _coerce_blocking_row(row: tuple[Any, ...]) -> tuple[str, int, str]:
    return (str(row[0]), int(row[1] or 0), str(row[2] or "0.0%"))


def _coerce_error_total_row(row: tuple[Any, ...]) -> tuple[str, int, int, str]:
    return (str(row[0]), int(row[1] or 0), int(row[2] or 0), str(row[3] or "0.0%"))


def _coerce_error_breakdown_row(row: tuple[Any, ...]) -> tuple[str, str, str, int, int, str]:
    return (
        str(row[0]),
        str(row[1]),
        str(row[2]),
        int(row[3] or 0),
        int(row[4] or 0),
        str(row[5] or "0.0%"),
    )
